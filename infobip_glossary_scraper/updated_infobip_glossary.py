import json
import logging
from docx import Document

# Set up logging
logging.basicConfig(filename="glossary.log", level=logging.INFO, 
                    format="%(asctime)s - %(levelname)s - %(message)s")

def save_to_word(data, filename):
    """Saves glossary data to a formatted Word document."""
    try:
        doc = Document()
        doc.add_heading("Infobip Glossary", level=1)

        for entry in data:
            term = entry.get("term", "N/A").strip()
            link = entry.get("link", "").strip()
            definition = entry.get("definition", "N/A").strip()
            header = entry.get("header", "").strip()
            paragraphs = entry.get("paragraphs", [])

            doc.add_heading(term, level=2)

            if link:
                doc.add_paragraph(f"More info: {link}", style="Intense Quote")

            if definition != "N/A":
                doc.add_paragraph(definition, style="Normal")

            if header:
                doc.add_paragraph(header, style="Heading 3")

            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            doc.add_paragraph("\n")

        doc.save(filename)
        logging.info(f"Glossary saved to Word document: {filename}")

    except Exception as e:
        logging.error(f"Error while writing Word document: {e}")
        raise

def json_to_word(json_filename, word_filename):
    """Reads glossary data from a JSON file and saves it as a Word document."""
    try:
        with open(json_filename, "r", encoding="utf-8") as file:
            data = json.load(file)

        if not isinstance(data, list):
            raise ValueError("Invalid JSON format: Expected a list of glossary entries")

        save_to_word(data, word_filename)
        logging.info(f"Successfully converted {json_filename} to {word_filename}")

    except json.JSONDecodeError:
        logging.error("Invalid JSON file format.")
    except FileNotFoundError:
        logging.error(f"File not found: {json_filename}")
    except Exception as e:
        logging.error(f"Error while converting JSON to Word: {e}")
        raise

if __name__ == "__main__":
    json_to_word("updated_infobip_glossary.json", "Infobip_Glossary.docx")
