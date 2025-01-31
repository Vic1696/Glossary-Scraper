from docx import Document
import logging

def save_to_word(data, filename):
    """Saves glossary data to a formatted Word document."""
    try:
        doc = Document()
        doc.add_heading("Twilio Glossary", level=1)

        for entry in data:
            term = entry.get("term", "N/A")
            link = entry.get("link", "N/A")
            definition = entry.get("definition", "N/A")
            header = entry.get("header", "N/A")
            paragraphs = entry.get("paragraphs", [])

            doc.add_heading(term, level=2)

            if link != "N/A":
                doc.add_paragraph(f"More info: {link}")

            doc.add_paragraph(definition, style="Intense Quote")

            if header and header != "N/A":
                doc.add_paragraph(header, style="Heading 3")

            for para in paragraphs:
                doc.add_paragraph(para)

            doc.add_paragraph("\n")

        doc.save(filename)
        logging.info(f"Glossary saved to Word document: {filename}")

    except Exception as e:
        logging.error(f"Error while writing Word document: {e}")
        raise
