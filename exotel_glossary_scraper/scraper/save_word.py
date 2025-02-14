from docx import Document
import logging

def save_to_word(data, filename):
    try:
        doc = Document()
        doc.add_heading("Glossary", level=1)

        for entry in data:
            term = entry.get("term", "N/A")
            link = entry.get("link", "N/A")
            definition = entry.get("definition", "N/A")
            title = entry.get("title", "N/A")
            sections = entry.get("sections", {})

            # Add Term Title
            doc.add_heading(term, level=2)

            # Add Definition
            doc.add_paragraph("Definition:", style="Heading 3")
            doc.add_paragraph(definition)

            # Add Link (if available)
            if link != "N/A":
                doc.add_paragraph(f"More info: {link}", style="Intense Quote")

            # Add Additional Sections
            for section_title, content in sections.items():
                if section_title and section_title != "N/A":
                    doc.add_heading(section_title, level=3)
                for paragraph in content:
                    doc.add_paragraph(paragraph)

            # Add spacing between terms
            doc.add_paragraph("\n" + "-" * 50 + "\n")

        doc.save(filename)
        logging.info(f"Glossary saved to Word document: {filename}")

    except Exception as e:
        logging.error(f"Error while writing Word document: {e}")
        raise
